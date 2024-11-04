import docx
from docx import Document

# Function to create a Word document with specific content
def createtext(filename="rezawordfile.docx"):
    # Initialize a new Word document
    doc = docx.Document()

    # Add a main heading to the document
    doc.add_heading("Boutique Software Developer Company", 0)

    # Add a paragraph with introductory text
    doc.add_paragraph("This is the first page of this word document")

    # Add a subheading for the introduction section
    doc.add_heading("Introduction", 1)

    # Add a paragraph under the Introduction heading
    doc.add_paragraph("Now we are hiring many people for the help desk job.")

    # Add another subheading for the company timetable
    doc.add_heading("Company time table", 2)

    # Insert an image into the document (image file should exist at specified path)
    doc.add_picture("timetable.jpg")

    # Save the document to the specified filename
    doc.save(filename)

    # Return the filename for reference in other functions
    return filename


# Function to print the text content of a specified Word document
def printtext(filename="rezawordfile.docx"):
    # Load the Word document for reading
    doc = Document(filename)

    # Print a header indicating the start of text output
    print("\nlist of texts: ")

    # Loop through each paragraph in the document and print its text
    for i in doc.paragraphs:
        print(i.text)


# Main function to create and print the document's contents
def main():
    # Create the document and get its filename
    filename = createtext()

    # Print the contents of the created document
    printtext(filename)


# Entry point: run the main function if this script is executed directly
if __name__ == "__main__":
    main()
