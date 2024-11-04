import unittest
import os
from docx import Document
from assessment3_part2 import createtext, printtext
from io import StringIO
from unittest.mock import patch

class Testfunctions(unittest.TestCase):

    def setUp(self):
        # Define the filename for the test document
        self.filename = "test_rezawordfile.docx"
    
    def tearDown(self):
        # Clean up by removing the test document after each test
        if os.path.exists(self.filename):
            os.remove(self.filename)

    def test_createtext(self):
        # Create the document and verify it was created
        createtext(self.filename)
        self.assertTrue(os.path.exists(self.filename))

        # Open the document and verify its contents
        doc = Document(self.filename)
        self.assertEqual(doc.paragraphs[0].text, "Boutique Software Developer Company")
        self.assertEqual(doc.paragraphs[1].text, "This is the first page of this word document")
        self.assertEqual(doc.paragraphs[2].text, "Introduction")
        self.assertEqual(doc.paragraphs[3].text, "Now we are hiring many people for the help desk job.")
        self.assertEqual(doc.paragraphs[4].text, "Company time table")

    @patch('sys.stdout', new_callable=StringIO)
    def test_printtext(self, mock_stdout):
        # Create the document for testing printtext
        createtext(self.filename)

        # Call printtext and capture its output
        printtext(self.filename)
        output = mock_stdout.getvalue()

        # Check that each expected paragraph is in the printed output
        expected_texts = [
            "Boutique Software Developer Company",
            "This is the first page of this word document",
            "Introduction",
            "Now we are hiring many people for the help desk job.",
            "Company time table"
        ]
        
        for expected_text in expected_texts:
            self.assertIn(expected_text, output)

if __name__ == "__main__":
    unittest.main()
